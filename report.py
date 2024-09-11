import re
from html import unescape
from fpdf import FPDF
from docx import Document
import os

# Function to remove HTML tags and decode HTML entities
def strip_html_tags(text):
    clean = re.compile('<.*?>')
    text = re.sub(clean, '', text)  # Remove HTML tags
    text = re.sub(r'\.list-item\s*\{\s*display\s*:\s*block\s*;\s*\}', '', text)
    # Remove <style> tags and their contents
    text = re.sub(r'<style\b[^>]*>(.*?)<\/style>', '', text, flags=re.DOTALL)
    # Remove other HTML tags
    text = re.sub(r'<.*?>', '', text)
    text = unescape(text)  # Decode HTML entities
    text = text.replace('’', "'")
    return text

class PDF(FPDF):
    def header(self):
        # Add a header if necessary
        pass

    def footer(self):
        # Add a footer if necessary
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title, 0, 1, 'L')
        self.ln(5)

    def chapter_body(self, body):
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, body)
        self.ln()

def generate_pdf(data, folder, filename):
    try:
        pdf = PDF()
        pdf.add_page()

        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="CVE Report", ln=True, align='C')

        pdf.set_font("Arial", '', 12)

        for key, value in data.items():
            if key == 'ai_response':
                continue
            if key == 'Exploits':
                continue
            if isinstance(value, list):
                value = ", ".join(value)
            pdf.multi_cell(0, 10, txt=f"{key}: {value}", align='L')  # Use multi_cell for wrapping text

        if 'ai_response' in data and data['ai_response'].strip():
            pdf.add_page()
            pdf.chapter_title("AI-Generated Summary")
            stripped_response = strip_html_tags(data['ai_response'])
            pdf.chapter_body(stripped_response)
        else:
            print("Warning: 'ai_response' is missing or empty.")

        pdf.output(os.path.join(folder, filename))
    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise

def generate_docx(data, folder, filename):
    try:
        doc = Document()
        doc.add_heading('CVE Report', 0)

        for key, value in data.items():
            if key == 'ai_response':
                continue
            if key == 'Exploits':
                continue
            if isinstance(value, list):
                value = ", ".join(value)
            doc.add_paragraph(f"{key}: {value}")
        
        if 'ai_response' in data and data['ai_response'].strip():
            doc.add_heading('AI Response:', level=1)
            stripped_response = strip_html_tags(data['ai_response'])
            doc.add_paragraph(stripped_response)  # Ensure it's plain text
        
        doc.save(os.path.join(folder, filename))
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        raise

def generate_markdown(data, folder, filename):
    try:
        md_content = "# CVE Report\n\n"
        
        for key, value in data.items():
            if key == 'ai_response':
                continue  # Skip 'ai_response' here to avoid duplication
            if key == 'Exploits':
                continue
            if isinstance(value, list):
                value = ", ".join(value)
            md_content += f"**{key}**: {value}\n\n"
        
        if 'ai_response' in data and data['ai_response'].strip():
            md_content += "## AI Response\n\n"
            stripped_response = strip_html_tags(data['ai_response'])
            md_content += f"{stripped_response}\n\n"

        with open(os.path.join(folder, filename), "w") as md_file:
            md_file.write(md_content)
    
    except Exception as e:
        print(f"Error generating Markdown: {e}")
        raise
